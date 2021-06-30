#!/usr/bin/env python3
# -*- coding: utf-8 -*-

'''
name:
cosmic_tool_3.py
email:
yinglj@asiainfo.com
title:
cosmic_tool_3.py for Linux
'''

import os
import sys
import openpyxl
from progressbar import *
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn

class CCosmic:
    def __init__(self):
        self.listProcessFiles = []
        self.columnNames = {}   #excel中的列名称
        self.cosmic_docx = None
        self.docx_file = None
        pass
                
    def processAllFiles(self):
        if not os.path.exists(sys.argv[1]):
            print('Excel template file({}) not exists.'.format(sys.argv[1]))
            return False
        # 新建空白文档
        self.cosmic_docx = Document()
        self.docx_file = sys.argv[1]

        for i in range(2, len(sys.argv)):
            self.listProcessFiles.append(sys.argv[i])
        for v in self.listProcessFiles:
            self.processOneFile(v)
        return True
            
    def processOneFile(self, cosmic_excel_file):
        if not os.path.exists(cosmic_excel_file):
            print('{} file not exist.'.format(cosmic_excel_file))
            return False
        
        wb = openpyxl.load_workbook(cosmic_excel_file)
        ws = wb.active

        START_ROW = 2           #从第二行开始解析
        FIRST_FUNCTION_COL = 2  #一级功能在第3列
        SECOND_FUNCTION_COL= 4  #二级功能在第5列
        SUB_PROCESS_COL = 5     #子过程描述在第6列

        FIRST_FUNCTION_TOPIC = 4  #docx中第4级Topic
        SECOND_FUNCTION_TOPIC= 5  #docx中第5级Topic

        #进度条功能
        i = 0
        widgets = ['Progress: ', FormatLabel('%(value)d of %(max)d '), Percentage(), ' ', Bar('#'),' ', Timer(), ' ', ETA(), ' ', FileTransferSpeed()]
        pbar = ProgressBar(widgets=widgets, maxval=ws.max_row).start()     #进度条功能
        pbar.update_interval = 1
        restart_numbering = True
        tableList = None
        sn = 1
        for row in ws.iter_rows(min_row=START_ROW, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
            if row[FIRST_FUNCTION_COL].value is not None:   #*一级功能, 该单元格有值则写入docx
                self.cosmic_docx.add_heading(row[FIRST_FUNCTION_COL].value, FIRST_FUNCTION_TOPIC)
            
            if row[SECOND_FUNCTION_COL].value is not None:  #*二级功能, 该单元格有值则写入docx
                self.cosmic_docx.add_heading(row[SECOND_FUNCTION_COL].value, SECOND_FUNCTION_TOPIC)
                sn = 1
            else:
                sn = sn + 1
            
            if row[SUB_PROCESS_COL].value is not None:  #*子过程描述, 该单元格有值则写入docx
                p0 = self.cosmic_docx.add_paragraph("{}. {}".format(sn, row[SUB_PROCESS_COL].value))

            #更新进度条进度
            i = i + 1
            pbar.update(i)
        self.cosmic_docx.save(self.docx_file)
        return True

def Usage(command):
    print("usage:" + command + "word_file excel_file ")
    print("example: " + command + "cosmic.docx cosmic.xlsx ")

if __name__ == '__main__':
    if len(sys.argv) == 1:
        Usage(sys.argv[0])
    else:
        try:
            client = CCosmic()
            client.processAllFiles()
            time.sleep(0)
        except KeyboardInterrupt as e:
            print(e)
        except IOError as e:
            print(e)
        except ValueError as e:
            print(e)
        finally:
            pass